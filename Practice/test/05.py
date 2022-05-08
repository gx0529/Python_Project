import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

guesttxt = open('guests.txt', 'r')
doc = docx.Document()
name = []
for guestname in guesttxt.readlines():
    if guestname != '\n':
        name.append(guestname)
for guest in name:
    guestname = guest.split('\n')[0]
    para1 = doc.add_paragraph('It would be a pleasure to have the company of')
    para1.style = 'Quote'
    para2 = doc.add_paragraph(guest)
    para2.style = 'Caption'
    para3 = doc.add_paragraph('at 11010 Memory Lake on the Evening of')
    para3.style = 'List'
    para4 = doc.add_paragraph('April 1st')
    para5 = doc.add_paragraph("at 7 o'clock")
    doc.add_page_break()
len = len(doc.paragraphs)
for i in range(0, 30):
    doc.paragraphs[i].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 全部居中
for j in range(0, 30, 2):
    doc.paragraphs[j].runs[0].bold = True
doc.save('guests-invatation.docx')